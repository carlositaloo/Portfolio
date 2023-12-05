import PyPDF2
from docx import Document
import os

def convert_pdfs_to_word(pdf_folder, output_word_file):
    # Criar um novo documento Word
    doc = Document()

    # Listar todos os arquivos PDF no diretório especificado
    for filename in os.listdir(pdf_folder):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(pdf_folder, filename)

            # Abrir o arquivo PDF
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)

                # Adicionar um título para cada PDF no documento Word
                doc.add_heading(filename, level=1)

                # Ler cada página do PDF
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)

                # Adicionar uma quebra de página no Word após cada PDF
                doc.add_page_break()

    # Salvar o documento Word
    doc.save(output_word_file)

# Caminho da pasta com os arquivos PDF
pdf_folder = 'PDF/'

# Nome do arquivo Word de saída
output_word_file = 'saida.docx'

convert_pdfs_to_word(pdf_folder, output_word_file)
